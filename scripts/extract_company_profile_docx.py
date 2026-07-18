#!/usr/bin/env python3
"""Extract the BATKO.DIGITAL.AI DOCX profile into Markdown and JSON."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "assets" / "company_profile" / "BATKO_DIGITAL_AI_Profil_2026_Rozsirena_Verze.docx"
OUT_MD = ROOT / "assets" / "company_profile" / "BATKO_DIGITAL_AI_Profil_2026_Rozsirena_Verze.md"
OUT_JSON = ROOT / "assets" / "company_profile" / "BATKO_DIGITAL_AI_Profil_2026_Rozsirena_Verze.json"


def clean(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    return text


def main() -> None:
    doc = Document(SRC)
    blocks: list[dict] = []
    md_lines: list[str] = [
        "# BATKO.DIGITAL.AI Profil 2026 - rozšířená verze",
        "",
        f"Zdrojový DOCX: `{SRC.name}`",
        "",
        "> Autoritativní profil pro šablony, nabídky, aukce, Codex a Cowork. Neupravovat ručně odhadem; změny dělat primárně ve zdrojovém DOCX nebo v tomto extraktu po ověření.",
        "",
    ]

    for p in doc.paragraphs:
        text = clean(p.text)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        level = 0
        m = re.match(r"Heading ([1-6])", style)
        if m:
            level = int(m.group(1))
            md_lines.append(f"{'#' * min(level + 1, 6)} {text}")
        elif style.lower().startswith("title"):
            md_lines.append(f"## {text}")
        else:
            md_lines.append(text)
        md_lines.append("")
        blocks.append({"type": "paragraph", "style": style, "text": text})

    for ti, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        blocks.append({"type": "table", "index": ti, "rows": rows})
        md_lines.append(f"## Tabulka {ti}")
        width = max(len(r) for r in rows)
        padded = [r + [""] * (width - len(r)) for r in rows]
        md_lines.append("| " + " | ".join(padded[0]) + " |")
        md_lines.append("| " + " | ".join(["---"] * width) + " |")
        for row in padded[1:]:
            md_lines.append("| " + " | ".join(row) + " |")
        md_lines.append("")

    data = {
        "source_docx": str(SRC.relative_to(ROOT)).replace("\\", "/"),
        "markdown": str(OUT_MD.relative_to(ROOT)).replace("\\", "/"),
        "blocks": blocks,
    }
    OUT_MD.write_text("\n".join(md_lines).rstrip() + "\n", encoding="utf-8")
    OUT_JSON.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"OK {OUT_MD}")
    print(f"OK {OUT_JSON}")
    print(f"paragraphs/tables: {len(blocks)}")


if __name__ == "__main__":
    main()
